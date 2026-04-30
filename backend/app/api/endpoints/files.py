# backend/app/api/endpoints/files.py
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
import os
import uuid
from app.core.sync_database import SessionLocal
from app.models.sql.file import File
from app.models.sql.user import User
from app.api.dependencies.auth import get_current_user
from app.services.rag_service import RAGService

router = APIRouter()

UPLOAD_DIR = "uploads"

def process_file_for_rag(file_id: int, file_path: str, original_filename: str, uploader_id: int, uploader_name: str):
    """后台任务：处理文件并添加到向量库"""
    db = SessionLocal()
    try:
        print(f"[RAG-后台] 开始处理文件 {file_id}: {original_filename}")
        
        # 提取带页码的文本内容
        pages_content = extract_text_with_pages(file_path, original_filename)
        total_pages = len(pages_content)
        total_chars = sum(len(p["text"]) for p in pages_content)
        print(f"[RAG-后台] 提取到 {total_pages} 页，共 {total_chars} 字符")
        
        # 初始化 RAG 服务并添加文档
        rag_service = RAGService()
        base_metadata = {
            "file_id": file_id,
            "filename": original_filename,
            "uploader_id": uploader_id,
            "uploader_name": uploader_name,
            "total_pages": total_pages
        }
        rag_service.add_document_with_pages(pages_content, base_metadata)
        
        # 更新数据库中的处理状态
        db_file = db.query(File).filter(File.id == file_id).first()
        if db_file:
            db_file.is_processed = True
            db.commit()
        
        print(f"[RAG-后台] 文件 {file_id} 处理完成")
    except Exception as e:
        print(f"[RAG-后台] 处理文件 {file_id} 失败: {e}")
        import traceback
        traceback.print_exc()
    finally:
        db.close()

def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    从文件中提取文本内容（简单版本，返回纯文本）
    支持：txt, md, pdf, docx 等格式
    """
    try:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext in ['.txt', '.md', '.text']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif ext == '.pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except ImportError:
                print("警告：PyPDF2 未安装，无法提取 PDF 内容")
                return f"[PDF 文件：{filename}]"
        
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
            except ImportError:
                print("警告：python-docx 未安装，无法提取 DOCX 内容")
                return f"[Word 文档：{filename}]"
        
        else:
            return f"[文件：{filename}]"
    
    except Exception as e:
        print(f"提取文件文本失败：{e}")
        return f"[文件读取失败：{filename}]"

def extract_text_with_pages(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """
    从文件中提取文本内容，返回带页码信息的结构化数据
    返回格式：[{"page": 1, "text": "..."}, {"page": 2, "text": "..."}]
    """
    try:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext in ['.txt', '.md', '.text']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return [{"page": 1, "text": content}]
        
        elif ext == '.pdf':
            try:
                import PyPDF2
                pages_content = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    total_pages = len(reader.pages)
                    print(f"[PDF提取] 总页数: {total_pages}")
                    
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = page.extract_text()
                        print(f"[PDF提取] 第{page_num}页原始文本长度: {len(text) if text else 0}")
                        if text and text.strip():
                            pages_content.append({
                                "page": page_num,
                                "text": text.strip()
                            })
                            print(f"[PDF提取] 第{page_num}页: {len(text)} 字符")
                        else:
                            print(f"[PDF提取] 第{page_num}页: 无文本内容（可能是图片或扫描版）")
                    
                    # 如果没有提取到任何文本，可能是扫描版PDF
                    if not pages_content:
                        print(f"[PDF提取] 警告：未提取到任何文本，该PDF可能是扫描版或图片型PDF")
                        # 尝试使用pdfplumber作为备选
                        try:
                            import pdfplumber
                            print(f"[PDF提取] 尝试使用pdfplumber提取...")
                            with pdfplumber.open(file_path) as pdf:
                                for page_num, page in enumerate(pdf.pages, start=1):
                                    text = page.extract_text()
                                    if text and text.strip():
                                        pages_content.append({
                                            "page": page_num,
                                            "text": text.strip()
                                        })
                                        print(f"[PDF提取-pdfplumber] 第{page_num}页: {len(text)} 字符")
                        except ImportError:
                            print("[PDF提取] pdfplumber未安装，跳过备选方案")
                
                return pages_content
            except ImportError:
                print("警告：PyPDF2 未安装，无法提取 PDF 内容")
                return [{"page": 1, "text": f"[PDF 文件：{filename}]"}]
        
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return [{"page": 1, "text": text}]
            except ImportError:
                print("警告：python-docx 未安装，无法提取 DOCX 内容")
                return [{"page": 1, "text": f"[Word 文档：{filename}]"}]
        
        else:
            return [{"page": 1, "text": f"[文件：{filename}]"}]
    
    except Exception as e:
        print(f"提取文件文本失败：{e}")
        return [{"page": 1, "text": f"[文件读取失败：{filename}]"}]

@router.post("/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(),
    current_user: User = Depends(get_current_user),
    is_knowledge_base: Optional[bool] = None
):
    """上传文件（异步处理向量索引）"""
    if not current_user.is_approved:
        raise HTTPException(status_code=403, detail="Account not approved")
    
    # 检查权限：只有教师和管理员可以上传到知识库
    if is_knowledge_base and current_user.role not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Only teachers and admins can upload to knowledge base")
    
    # 确保上传目录存在
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    # 生成唯一文件名
    file_ext = os.path.splitext(file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    # 保存文件
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    # 保存到数据库
    db = SessionLocal()
    db_file = File(
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=len(content),
        file_type=file.content_type,
        uploader_id=current_user.id,
        is_knowledge_base=is_knowledge_base or False,
        is_processed=False  # 初始状态为未处理
    )
    db.add(db_file)
    db.commit()
    db.refresh(db_file)
    file_id = db_file.id
    db.close()
    
    # 如果是知识库文件，添加后台任务处理向量索引
    if is_knowledge_base:
        background_tasks.add_task(
            process_file_for_rag,
            file_id,
            file_path,
            file.filename,
            current_user.id,
            current_user.username
        )
        print(f"[上传] 文件 {file_id} 已保存，向量索引正在后台处理...")
    
    return {
        "id": file_id,
        "filename": file.filename,
        "size": len(content),
        "path": file_path,
        "is_processing": is_knowledge_base or False
    }

@router.get("/list")
def list_files(
    current_user: User = Depends(get_current_user)
):
    """获取文件列表"""
    db = SessionLocal()
    files = db.query(File).filter(File.uploader_id == current_user.id).all()
    db.close()
    return files

@router.get("/download/{file_id}")
def download_file(
    file_id: int,
    token: str,  # 从 query 参数获取 token
):
    """下载文件"""
    from fastapi.responses import FileResponse, Response
    from jose import JWTError, jwt
    from app.core.config import settings
    
    # 验证 token
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Not authenticated")
    except JWTError:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    db = SessionLocal()
    file = db.query(File).filter(File.id == file_id).first()
    db.close()
    
    if not file:
        raise HTTPException(status_code=404, detail="File not found")
    
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")
    
    # 判断是否应该在线预览（inline）而不是下载（attachment）
    preview_types = ['application/pdf', 'image/jpeg', 'image/png', 'image/gif', 'image/webp', 'text/plain', 'text/markdown', 'application/json', 'text/csv']
    file_type = file.file_type or "application/octet-stream"  # 如果 file_type 为 None，使用默认值
    media_type = file_type if file_type in preview_types else "application/octet-stream"
    
    # 检查文件是否存在
    if not os.path.exists(file.file_path):
        db.close()
        raise HTTPException(status_code=404, detail="File not found on disk")
    
    # 创建 FileResponse 并添加 CORS 头
    try:
        import urllib.parse
        from fastapi.responses import StreamingResponse
        
        # 读取文件内容
        with open(file.file_path, 'rb') as f:
            file_content = f.read()
        
        # 判断是否应该在线预览（inline）而不是下载（attachment）
        if file_type in preview_types:
            # 在线预览类型使用 inline
            disposition = f'inline; filename="{urllib.parse.quote(file.original_filename)}"'
        else:
            # 其他类型使用 attachment 下载
            disposition = f'attachment; filename="{urllib.parse.quote(file.original_filename)}"'
        
        # 创建响应
        response = Response(
            content=file_content,
            media_type=media_type
        )
        
        # 添加 CORS 响应头
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Content-Disposition"] = disposition
        response.headers["Content-Length"] = str(len(file_content))
        
        db.close()
        return response
    except Exception as e:
        db.close()
        print(f"文件响应创建失败：{e}")
        raise HTTPException(status_code=500, detail=f"文件处理失败：{str(e)}")

@router.delete("/{file_id}")
def delete_file(
    file_id: int,
    current_user: User = Depends(get_current_user)
):
    """删除文件（管理员和教师）"""
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Admin or Teacher only")
    
    db = SessionLocal()
    file = db.query(File).filter(File.id == file_id).first()
    
    if not file:
        db.close()
        raise HTTPException(status_code=404, detail="File not found")
    
    # 删除磁盘上的文件
    if os.path.exists(file.file_path):
        try:
            os.remove(file.file_path)
            print(f"[File] 已删除磁盘文件: {file.file_path}")
        except Exception as e:
            print(f"[File] 删除磁盘文件失败: {e}")
    
    # 如果是知识库文件，删除 RAG 向量索引
    if file.is_knowledge_base:
        try:
            rag_service = RAGService()
            rag_service.delete_document(file_id)
            print(f"[RAG] 已删除文件 {file_id} 的向量索引")
        except Exception as e:
            print(f"[RAG] 删除向量索引失败: {e}")
            # 向量索引删除失败不阻止文件删除
    
    # 删除数据库记录
    db.delete(file)
    db.commit()
    db.close()
    
    return {"msg": "File deleted", "file_id": file_id}

@router.get("/knowledge-base")
def list_knowledge_base_files(
    current_user: User = Depends(get_current_user),
    page: int = 1,
    page_size: int = 10
):
    """获取知识库文件列表（支持分页）"""
    if not current_user.is_approved:
        raise HTTPException(status_code=403, detail="Account not approved")

    db = SessionLocal()
    try:
        # 获取总记录数
        total = db.query(File).count()

        # 分页查询
        offset = (page - 1) * page_size
        files = db.query(File).offset(offset).limit(page_size).all()

        # 转换为字典格式，包含上传者信息
        result = []
        for file in files:
            file_dict = {
                "id": file.id,
                "filename": file.original_filename,
                "original_filename": file.original_filename,
                "file_size": file.file_size,
                "file_type": file.file_type,
                "created_at": file.created_at,
                "uploader_id": file.uploader_id,
                "uploader_name": None
            }

            # 获取上传者信息
            uploader = db.query(User).filter(User.id == file.uploader_id).first()
            if uploader:
                file_dict["uploader_name"] = uploader.username

            result.append(file_dict)

        return {
            "items": result,
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": (total + page_size - 1) // page_size
        }
    finally:
        db.close()

@router.get("/debug/chroma")
def debug_chroma(
    current_user: User = Depends(get_current_user)
):
    """调试接口：查看向量库中的数据"""
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Admin or Teacher only")
    
    try:
        rag_service = RAGService()
        
        # 获取所有数据
        all_data = rag_service.collection.get()
        
        result = {
            "total_count": len(all_data['ids']) if all_data['ids'] else 0,
            "ids": all_data['ids'][:20] if all_data['ids'] else [],
            "documents": [],
            "metadatas": all_data['metadatas'][:20] if all_data['metadatas'] else []
        }
        
        # 显示文档内容（截取前200字符）
        if all_data['documents']:
            for doc in all_data['documents'][:20]:
                result["documents"].append(doc[:200] + "..." if len(doc) > 200 else doc)
        
        return result
    except Exception as e:
        return {"error": str(e)}

@router.delete("/debug/chroma")
def clear_chroma(
    current_user: User = Depends(get_current_user)
):
    """调试接口：清空向量库"""
    if current_user.role not in ["admin"]:
        raise HTTPException(status_code=403, detail="Admin only")
    
    try:
        rag_service = RAGService()
        
        # 获取所有ID
        all_data = rag_service.collection.get()
        total_count = len(all_data['ids']) if all_data['ids'] else 0
        
        if total_count > 0:
            # 删除所有数据
            rag_service.collection.delete(ids=all_data['ids'])
            print(f"[RAG] 已清空向量库，删除了 {total_count} 条记录")
        
        return {
            "msg": "ChromaDB cleared",
            "deleted_count": total_count
        }
    except Exception as e:
        return {"error": str(e)}

@router.post("/debug/rebuild-index/{file_id}")
def rebuild_file_index(
    file_id: int,
    current_user: User = Depends(get_current_user)
):
    """调试接口：重新构建指定文件的向量索引"""
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Admin or Teacher only")
    
    db = SessionLocal()
    file = db.query(File).filter(File.id == file_id).first()
    
    if not file:
        db.close()
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        # 先删除旧的索引
        rag_service = RAGService()
        rag_service.delete_document(file_id)
        print(f"[RAG] 已删除旧索引: file_id={file_id}")
        
        # 重新提取文本
        text_content = extract_text_from_file(file.file_path, file.original_filename)
        print(f"[RAG] 重新提取文本长度: {len(text_content)}")
        
        # 重新添加到向量库
        metadata = {
            "file_id": file.id,
            "filename": file.original_filename,
            "uploader_id": file.uploader_id,
            "uploader_name": current_user.username
        }
        rag_service.add_document(text_content, metadata)
        
        db.close()
        return {
            "msg": "Index rebuilt successfully",
            "file_id": file_id,
            "filename": file.original_filename,
            "text_length": len(text_content)
        }
    except Exception as e:
        db.close()
        return {"error": str(e)}
